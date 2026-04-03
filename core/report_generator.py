from __future__ import annotations

from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from functools import lru_cache
from io import BytesIO
import json
import textwrap
from typing import Dict

from core.llm_engine import get_llm_response_high_quality


REPORT_SECTION_ORDER = [
    "Report Objective",
    "Document Profile",
    "Executive Summary",
    "Overall Risk Rating",
    "Critical Issues",
    "Key Obligations",
    "Missing Protections / Negotiation Gaps",
    "Recommended Actions",
    "Conclusion",
    "Appendix - Supporting Notes",
]


@dataclass(frozen=True)
class ReportOptions:
    tone: str = "Formal"
    structure: str = "Detailed"
    focus_area: str = "Balanced"


def _empty_report_sections() -> Dict[str, str]:
    return {section: "" for section in REPORT_SECTION_ORDER}


def _is_placeholder_text(text: str) -> bool:
    value = " ".join((text or "").strip(" -.:").lower().split())
    if not value:
        return True
    placeholders = {
        "not stated",
        "not available",
        "none stated",
        "n/a",
        "na",
        "unknown",
        "tbd",
        "not generated",
        "auto-detected from uploaded analysis",
        "auto-detected from uploaded contract analysis",
    }
    return value in placeholders or value.startswith("not generated because") or value.startswith("not generated due to")


def _complete_sentence(text: str) -> str:
    value = " ".join((text or "").strip().split())
    if not value:
        return ""
    trailing_terms = (" and", " or", " to", " for", " with", " without", " of", " the", " a", " an")
    lowered = value.lower().rstrip(". ")
    if value.endswith("...") or value.endswith("..") or any(lowered.endswith(term) for term in trailing_terms):
        value = value.rstrip(". ,;:-")
        words = value.split()
        while words:
            candidate = " ".join(words)
            lowered_candidate = candidate.lower()
            if not any(lowered_candidate.endswith(term) for term in trailing_terms):
                value = candidate
                break
            words.pop()
        value = value.strip(" ,;:-")
    if not value:
        return ""
    if value[-1] not in ".!?":
        value += "."
    return value


def _candidate_sentences(text: str, limit: int = 6) -> list[str]:
    normalized = " ".join((text or "").strip().split())
    if not normalized or _is_placeholder_text(normalized):
        return []
    parts = textwrap.wrap(normalized, width=400, break_long_words=False, break_on_hyphens=False)
    source = " ".join(parts)
    sentences = []
    seen = set()
    for raw in source.replace("\n", " ").split(". "):
        cleaned = _complete_sentence(raw.lstrip("-").strip())
        if not cleaned or _is_placeholder_text(cleaned):
            continue
        key = cleaned.lower()
        if key in seen:
            continue
        seen.add(key)
        sentences.append(cleaned)
        if len(sentences) >= limit:
            break
    return sentences


def _split_report_sections(text: str) -> Dict[str, str]:
    sections = _empty_report_sections()
    aliases = {
        "report objective": "Report Objective",
        "document profile": "Document Profile",
        "executive summary": "Executive Summary",
        "plain summary": "Executive Summary",
        "summary": "Executive Summary",
        "overall risk rating": "Overall Risk Rating",
        "overall risk score": "Overall Risk Rating",
        "risk score": "Overall Risk Rating",
        "critical issues": "Critical Issues",
        "risk analysis": "Critical Issues",
        "risks": "Critical Issues",
        "key obligations": "Key Obligations",
        "key clauses": "Key Obligations",
        "key insights": "Key Obligations",
        "missing protections / negotiation gaps": "Missing Protections / Negotiation Gaps",
        "missing / weak points": "Missing Protections / Negotiation Gaps",
        "missing / weak": "Missing Protections / Negotiation Gaps",
        "recommended actions": "Recommended Actions",
        "recommendations": "Recommended Actions",
        "conclusion": "Conclusion",
        "appendix - supporting notes": "Appendix - Supporting Notes",
        "appendix": "Appendix - Supporting Notes",
    }
    current = None
    for raw in (text or "").splitlines():
        line = raw.strip()
        if not line:
            continue
        lowered = line.rstrip(":").lower()
        mapped = aliases.get(lowered)
        if mapped:
            current = mapped
            continue
        if current:
            sections[current] += raw.strip() + "\n"
    return {key: value.strip() for key, value in sections.items()}


def _build_fallback_profile(final_output: str) -> str:
    analysis_sections = _split_report_sections(final_output or "")
    summary = analysis_sections.get("Executive Summary") or "Summary generated from the available analysis."
    risk_rating = analysis_sections.get("Overall Risk Rating") or "Medium"
    return "\n".join(
        [
            "- Document type: Auto-detected from uploaded contract analysis",
            "- Review scope: Contract obligations, risk exposure, and weak protections",
            f"- Current assessment basis: {summary}",
            f"- Provisional risk rating: {risk_rating}",
        ]
    )


def _derive_report_title(sections: Dict[str, str]) -> str:
    profile_points = _section_points(sections.get("Document Profile", ""))
    for point in profile_points:
        lowered = point.lower()
        if lowered.startswith("document type:"):
            value = point.split(":", 1)[1].strip()
            if value:
                return f"{value} Review Report"
    obligations = _section_points(sections.get("Key Obligations", ""))
    if obligations and obligations[0] != "No supported detail was available for this section.":
        short_title = textwrap.shorten(obligations[0], width=54, placeholder="...")
        return f"Contract Review Report: {short_title}"
    return "Contract Review Report"


def _fallback_report(final_output: str, agent_outputs: Dict[str, str], options: ReportOptions) -> str:
    sections = _split_report_sections(final_output or "")
    recommendation_lines = []
    appendix_lines = []
    for role, role_text in agent_outputs.items():
        trimmed = (role_text or "").strip()
        if not trimmed:
            continue
        appendix_lines.append(f"{role.title()} Agent:")
        appendix_lines.append(trimmed.strip())
        appendix_lines.append("")
        for line in trimmed.splitlines():
            stripped = line.strip()
            if stripped.startswith("-") and len(recommendation_lines) < 4:
                recommendation_lines.append(stripped)

    executive_summary = sections.get("Executive Summary") or "This report was prepared from the available contract analysis and agent findings."
    overall_risk = sections.get("Overall Risk Rating") or "Medium"
    fallback_sentences = _candidate_sentences(final_output, limit=6)
    critical_issues = sections.get("Critical Issues") or "\n".join(f"- {item}" for item in fallback_sentences[1:3]) or "- No material issue was confirmed from the available analysis."
    key_obligations = sections.get("Key Obligations") or "\n".join(f"- {item}" for item in fallback_sentences[:2]) or "- No obligation could be cleanly extracted from the available analysis."
    missing_gaps = sections.get("Missing Protections / Negotiation Gaps") or "\n".join(
        f"- {item}" for item in fallback_sentences if any(term in item.lower() for term in ("missing", "unclear", "not specify", "not defined", "weak", "silent"))
    ) or "- Review the contract for missing protections and negotiation gaps."
    recommended_actions = "\n".join(recommendation_lines) if recommendation_lines else "- Review the contract against internal legal and business requirements before approval."
    appendix = "\n".join(line for line in appendix_lines if line.strip()) or "Supporting agent notes were not available."

    return (
        "Report Objective:\n"
        "Provide a formal contract review highlighting material obligations, risk exposure, missing protections, and recommended follow-up actions."
        "\n\nDocument Profile:\n"
        + _build_fallback_profile(final_output)
        + "\n\nExecutive Summary:\n"
        + executive_summary
        + "\n\nOverall Risk Rating:\n"
        + overall_risk
        + "\n\nCritical Issues:\n"
        + critical_issues
        + "\n\nKey Obligations:\n"
        + key_obligations
        + "\n\nMissing Protections / Negotiation Gaps:\n"
        + missing_gaps
        + "\n\nRecommended Actions:\n"
        + recommended_actions
        + "\n\nConclusion:\n"
        + "The agreement should proceed only after the highlighted issues and recommended actions are reviewed by the appropriate stakeholders."
        + "\n\nAppendix - Supporting Notes:\n"
        + appendix
    )


def _build_report_prompt(final_output: str, agent_outputs: Dict[str, str], options: ReportOptions) -> str:
    agent_notes = "\n\n".join(f"{role.title()} Agent:\n{text}" for role, text in agent_outputs.items()) or "No agent notes were available."
    return f"""
You are preparing a formal contract review report for a professional audience.

Report customization:
- Tone: {options.tone}
- Structure: {options.structure}
- Focus area: {options.focus_area}

Instructions:
- Use only the provided analysis inputs.
- Make the report clearly different from raw analysis output.
- Write as a polished formal report suitable for managers, legal reviewers, or clients.
- Remove redundancy and avoid repeating the same points in multiple sections.
- Be thorough and comprehensive 2014 include ALL material clauses, obligations, risks, and findings.
- Each section should contain detailed, specific information from the contract.
- Use exact figures, dates, amounts, and references from the document wherever possible.
- The report should be formal, professional, and suitable for executive review.
- Keep the executive summary decision-oriented.
- Cover the full substance of the uploaded document as reflected in the analysis and agent outputs.
- Include all material clauses, obligations, risks, missing protections, and follow-up actions that can be supported by the provided inputs.
- Do not give a shallow summary; provide detailed but disciplined coverage in each section.
- The overall risk rating must be a short label such as Low, Medium, High, or Medium-High with one supporting sentence.
- Document profile should summarize core contract facts visible from the analysis.
- Critical issues should cover only the most material concerns.
- Key obligations should focus on the main duties, commitments, and operational responsibilities.
- Recommended actions must be specific and actionable.
- The conclusion should read like a formal closing assessment.
- Appendix should briefly summarize supporting agent observations rather than repeating the full report.
- Do not add unsupported facts.

Return exactly these sections:
Report Objective:
...

Document Profile:
- ...

Executive Summary:
...

Overall Risk Rating:
...

Critical Issues:
- ...

Key Obligations:
- ...

Missing Protections / Negotiation Gaps:
- ...

Recommended Actions:
- ...

Conclusion:
...

Appendix - Supporting Notes:
- ...

Base analysis:
{final_output or "No source analysis was available."}

Agent outputs:
{agent_notes}
""".strip()


def _section_lines(body: str) -> list[str]:
    lines = [line.strip() for line in (body or "").splitlines() if line.strip()]
    cleaned = [_complete_sentence(line) for line in lines if not _is_placeholder_text(line)]
    return [line for line in cleaned if line] or ["No supported detail was available for this section."]


def _section_points(body: str) -> list[str]:
    points = []
    for line in _section_lines(body):
        cleaned = line.lstrip("-").strip()
        if cleaned:
            points.append(cleaned)
    return points or ["No supported detail was available for this section."]


def _fill_report_sections(sections: Dict[str, str], final_output: str, agent_outputs: Dict[str, str]) -> Dict[str, str]:
    resolved = dict(sections)
    analysis_sentences = _candidate_sentences(final_output, limit=8)
    agent_sentences = _candidate_sentences("\n".join(text for text in agent_outputs.values() if text), limit=8)
    combined = analysis_sentences + [item for item in agent_sentences if item.lower() not in {s.lower() for s in analysis_sentences}]
    risk_points = [item for item in combined if any(term in item.lower() for term in ("risk", "liability", "termination", "indemn", "breach", "penalty", "missing", "unclear", "weak"))]
    missing_points = [item for item in combined if any(term in item.lower() for term in ("missing", "unclear", "not specify", "not defined", "weak", "silent", "absent"))]

    fallback_map = {
        "Report Objective": "Provide a contract review highlighting obligations, risks, weak protections, and recommended next steps.",
        "Document Profile": "\n".join([
            "- Document type: Auto-detected from uploaded contract analysis.",
            "- Review scope: Contract obligations, risk exposure, and weak protections.",
            f"- Current assessment basis: {(analysis_sentences[0] if analysis_sentences else 'Summary generated from the available analysis.')}",
        ]),
        "Executive Summary": " ".join((analysis_sentences or risk_points or ["Summary generated from the available analysis."])[:2]),
        "Overall Risk Rating": "Medium.",
        "Critical Issues": "\n".join(f"- {item}" for item in (risk_points[:3] or combined[:2])) or "- No material issue was confirmed from the available analysis.",
        "Key Obligations": "\n".join(f"- {item}" for item in (analysis_sentences[:3] or combined[:3])) or "- No obligation could be cleanly extracted from the available analysis.",
        "Missing Protections / Negotiation Gaps": "\n".join(f"- {item}" for item in (missing_points[:3] or risk_points[:2])) or "- Review the contract for missing protections and negotiation gaps.",
        "Recommended Actions": "\n".join(
            f"- Review and address: {item.rstrip('.')}."
            for item in (missing_points[:2] or risk_points[:2] or combined[:2])
        ) or "- Review the highlighted findings and confirm whether additional protections are needed.",
        "Conclusion": analysis_sentences[0] if analysis_sentences else "The agreement should proceed only after the highlighted issues are reviewed by the appropriate stakeholders.",
        "Appendix - Supporting Notes": "\n".join(f"- {item}" for item in (agent_sentences[:4] or combined[:4])) or "- Supporting notes were limited in the available analysis.",
    }

    for title in REPORT_SECTION_ORDER:
        current = (resolved.get(title) or "").strip()
        if not current or _is_placeholder_text(current):
            resolved[title] = fallback_map[title]

    resolved["Executive Summary"] = _complete_sentence(resolved["Executive Summary"]) or fallback_map["Executive Summary"]
    resolved["Report Objective"] = _complete_sentence(resolved["Report Objective"]) or fallback_map["Report Objective"]
    resolved["Conclusion"] = _complete_sentence(resolved["Conclusion"]) or fallback_map["Conclusion"]
    resolved["Overall Risk Rating"] = _complete_sentence(resolved["Overall Risk Rating"]) or "Medium."
    return resolved


def _compose_report_text(sections: Dict[str, str]) -> str:
    parts = []
    for title in REPORT_SECTION_ORDER:
        parts.append(f"{title}:")
        parts.append((sections.get(title) or "").strip())
        parts.append("")
    return "\n".join(parts).strip()


@lru_cache(maxsize=64)
def _generate_report_content_cached(final_output: str, agent_outputs_json: str, options_json: str) -> tuple[str, Dict[str, str]]:
    agent_outputs = json.loads(agent_outputs_json) if agent_outputs_json else {}
    options = ReportOptions(**(json.loads(options_json) if options_json else {}))
    prompt = _build_report_prompt(final_output, agent_outputs, options)
    try:
        report_text = get_llm_response_high_quality(
            prompt,
            num_predict=max(1500, int(os.getenv("CLAUSEAI_REPORT_NUM_PREDICT", "1500"))),
        )
    except Exception:
        report_text = _fallback_report(final_output, agent_outputs, options)

    sections = _split_report_sections(report_text)
    if not any(sections.values()):
        report_text = _fallback_report(final_output, agent_outputs, options)
        sections = _split_report_sections(report_text)
    sections = _fill_report_sections(sections, final_output, agent_outputs)
    report_text = _compose_report_text(sections)
    return report_text.strip(), sections


def generate_report(final_output: str, agent_outputs: Dict[str, str], options: ReportOptions) -> Dict[str, object]:
    agent_outputs_json = json.dumps(agent_outputs or {}, sort_keys=True, ensure_ascii=True, separators=(",", ":"))
    options_json = json.dumps(asdict(options), sort_keys=True, ensure_ascii=True, separators=(",", ":"))
    report_text, sections = _generate_report_content_cached(final_output, agent_outputs_json, options_json)
    metadata = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "options": asdict(options),
        "risk_score": sections.get("Overall Risk Rating", "").strip() or "Medium",
        "report_title": _derive_report_title(sections),
    }
    return {"report_text": report_text.strip(), "sections": sections, "metadata": metadata}


@lru_cache(maxsize=64)
def _build_report_docx_bytes_cached(report_text: str, metadata_json: str) -> bytes:
    metadata = json.loads(metadata_json) if metadata_json else None
    return _build_report_docx_bytes_uncached(report_text, metadata)


def build_report_docx_bytes(report_text: str, metadata: Dict[str, object] | None = None) -> bytes:
    metadata_json = json.dumps(metadata or {}, sort_keys=True, ensure_ascii=True, separators=(",", ":"))
    return _build_report_docx_bytes_cached(report_text, metadata_json)


def _build_report_docx_bytes_uncached(report_text: str, metadata: Dict[str, object] | None = None) -> bytes:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    sections = _split_report_sections(report_text)
    report_title = (
        str(metadata.get("report_title", "")).strip()
        if metadata else ""
    ) or _derive_report_title(sections)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(report_title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(23, 44, 79)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Generated by ClauseAI")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(96, 108, 128)

    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    label_map = [
        ("Generated", metadata.get("generated_at", "Unknown") if metadata else "Unknown"),
        ("Risk Rating", metadata.get("risk_score", "Medium") if metadata else "Medium"),
    ]
    for row_idx, (label, value) in enumerate(label_map):
        row = table.rows[row_idx]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        for cell_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    if cell_idx == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(23, 44, 79)

    document.add_paragraph("")
    if metadata:
        pass
    for title in REPORT_SECTION_ORDER:
        body = sections.get(title) or "No supported detail was available for this section."
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
        run = heading.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(26, 55, 94)

        if title in {"Document Profile", "Critical Issues", "Key Obligations", "Missing Protections / Negotiation Gaps", "Recommended Actions", "Appendix - Supporting Notes"}:
            for point in _section_points(body):
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.add_run(point)
        else:
            for line in _section_lines(body):
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.add_run(line)
    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


@lru_cache(maxsize=64)
def _build_report_pdf_bytes_cached(report_text: str) -> bytes:
    return _build_report_pdf_bytes_uncached(report_text)


def build_report_pdf_bytes(report_text: str) -> bytes:
    return _build_report_pdf_bytes_cached(report_text)


def _build_report_pdf_bytes_uncached(report_text: str) -> bytes:
    try:
        return _build_report_pdf_bytes_reportlab(report_text)
    except Exception:
        return _build_report_pdf_bytes_basic(report_text)


def _build_report_pdf_bytes_reportlab(report_text: str) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    sections = _split_report_sections(report_text)
    report_title = _derive_report_title(sections)
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    risk_rating = sections.get("Overall Risk Rating", "").strip() or "Medium"
    report_type = report_title.replace(" Review Report", "").strip() or "Contract Review"

    navy = colors.HexColor("#1a375e")
    slate = colors.HexColor("#5f6f86")
    border = colors.HexColor("#c7d2e3")
    text = colors.HexColor("#1f2937")
    muted_fill = colors.HexColor("#eef3f9")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ClauseAITitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        textColor=navy,
        spaceAfter=10,
    )
    body_style = ParagraphStyle(
        "ClauseAIBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        alignment=TA_JUSTIFY,
        textColor=text,
        spaceAfter=6,
    )
    bullet_style = ParagraphStyle(
        "ClauseAIBullet",
        parent=body_style,
        leftIndent=10,
        firstLineIndent=0,
        bulletIndent=0,
        spaceAfter=3,
    )
    heading_style = ParagraphStyle(
        "ClauseAIHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=16,
        textColor=navy,
        spaceBefore=10,
        spaceAfter=5,
    )
    meta_style = ParagraphStyle(
        "ClauseAIMeta",
        parent=body_style,
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        textColor=text,
        alignment=TA_CENTER,
        spaceAfter=0,
    )
    header_left_style = ParagraphStyle(
        "ClauseAIHeaderLeft",
        parent=body_style,
        fontName="Helvetica-Bold",
        fontSize=10,
        textColor=navy,
        spaceAfter=0,
    )
    header_right_style = ParagraphStyle(
        "ClauseAIHeaderRight",
        parent=body_style,
        fontName="Helvetica-Bold",
        fontSize=9.5,
        textColor=slate,
        alignment=TA_CENTER,
        spaceAfter=0,
    )

    def _header_footer(canvas, doc):
        canvas.saveState()
        width, height = letter
        margin = 0.55 * inch
        canvas.setStrokeColor(border)
        canvas.setLineWidth(1.5)
        canvas.rect(margin, margin, width - (2 * margin), height - (2 * margin))
        canvas.setFont("Helvetica-Bold", 10)
        canvas.setFillColor(navy)
        canvas.drawString(margin + 8, height - margin + 10, "ClauseAI Contract Review Report")
        canvas.setFont("Helvetica-Bold", 9)
        canvas.setFillColor(slate)
        confidential = "Confidential"
        confidential_width = canvas.stringWidth(confidential, "Helvetica-Bold", 9)
        canvas.drawString(width - margin - confidential_width - 8, height - margin + 10, confidential)
        canvas.setStrokeColor(border)
        canvas.line(margin + 8, height - margin + 4, width - margin - 8, height - margin + 4)
        footer = f"Page {canvas.getPageNumber()}"
        footer_width = canvas.stringWidth(footer, "Helvetica", 9)
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(slate)
        canvas.drawString((width - footer_width) / 2, margin - 18, footer)
        canvas.restoreState()

    story = [
        Paragraph(report_title, title_style),
        Paragraph("Confidential review prepared by ClauseAI", meta_style),
        Spacer(1, 0.12 * inch),
    ]

    metadata_table = Table(
        [
            [
                Paragraph("<b>Generation Date</b>", meta_style),
                Paragraph("<b>Risk Rating</b>", meta_style),
                Paragraph("<b>Report Type</b>", meta_style),
            ],
            [
                Paragraph(generated_at, meta_style),
                Paragraph(risk_rating, meta_style),
                Paragraph(report_type, meta_style),
            ],
        ],
        colWidths=[2.0 * inch, 1.5 * inch, 2.3 * inch],
        hAlign="CENTER",
    )
    metadata_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), muted_fill),
                ("TEXTCOLOR", (0, 0), (-1, -1), text),
                ("GRID", (0, 0), (-1, -1), 0.75, border),
                ("BOX", (0, 0), (-1, -1), 0.75, border),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([metadata_table, Spacer(1, 0.18 * inch)])

    bullet_sections = {
        "Document Profile",
        "Critical Issues",
        "Key Obligations",
        "Missing Protections / Negotiation Gaps",
        "Recommended Actions",
        "Appendix - Supporting Notes",
    }

    for title in REPORT_SECTION_ORDER:
        body = sections.get(title) or "No supported detail was available for this section."
        story.append(Paragraph(title, heading_style))
        story.append(HRFlowable(width="100%", thickness=1.5, color=border, spaceBefore=0, spaceAfter=7))
        if title in bullet_sections:
            items = [ListItem(Paragraph(point, bullet_style)) for point in _section_points(body)]
            story.append(
                ListFlowable(
                    items,
                    bulletType="bullet",
                    leftIndent=14,
                    bulletFontName="Helvetica",
                    bulletFontSize=10,
                )
            )
        else:
            for line in _section_lines(body):
                story.append(Paragraph(line, body_style))
        story.append(Spacer(1, 0.08 * inch))

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        topMargin=0.95 * inch,
        bottomMargin=0.85 * inch,
        leftMargin=0.78 * inch,
        rightMargin=0.78 * inch,
        title=report_title,
        author="ClauseAI",
    )
    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
    return buf.getvalue()


def _build_report_pdf_bytes_basic(report_text: str) -> bytes:
    page_width = 612
    page_height = 792
    margin = 50
    line_height = 15
    start_y = page_height - margin
    wrap_width = 84
    max_lines_per_page = max(1, int((page_height - (2 * margin)) / line_height))

    def _safe_pdf_text(line: str) -> str:
        normalized = line.encode("latin-1", "replace").decode("latin-1")
        return normalized.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    sections = _split_report_sections(report_text)
    logical_lines = [_derive_report_title(sections), "Generated by ClauseAI", ""]
    for title in REPORT_SECTION_ORDER:
        logical_lines.append(title.upper())
        body = sections.get(title) or "No supported detail was available for this section."
        for line in _section_lines(body):
            wrapped = textwrap.wrap(line.lstrip("- ").strip(), width=wrap_width) if line.strip() else [""]
            if line.strip().startswith("-"):
                if wrapped:
                    logical_lines.append(f"- {wrapped[0]}")
                    for continuation in wrapped[1:]:
                        logical_lines.append(f"  {continuation}")
                else:
                    logical_lines.append("- No supported detail was available for this section.")
            else:
                logical_lines.extend(wrapped or [""])
        logical_lines.append("")

    pages = [logical_lines[i : i + max_lines_per_page] for i in range(0, len(logical_lines), max_lines_per_page)]
    objects = {1: b"<< /Type /Catalog /Pages 2 0 R >>", 2: None, 3: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"}
    page_ids = []
    next_id = 4
    for page_lines in pages:
        page_id = next_id
        content_id = next_id + 1
        page_ids.append(page_id)
        next_id += 2
        stream_lines = ["BT", "/F1 11 Tf", f"{line_height} TL", f"{margin} {start_y} Td"]
        for idx, line in enumerate(page_lines):
            stream_lines.append(f"{'' if idx == 0 else 'T* '}({_safe_pdf_text(line)}) Tj")
        stream_lines.append("ET")
        stream_data = "\n".join(stream_lines).encode("latin-1", "replace")
        objects[content_id] = b"<< /Length " + str(len(stream_data)).encode("ascii") + b" >>\nstream\n" + stream_data + b"\nendstream"
        objects[page_id] = (
            b"<< /Type /Page /Parent 2 0 R "
            + f"/MediaBox [0 0 {page_width} {page_height}] ".encode("ascii")
            + b"/Resources << /Font << /F1 3 0 R >> >> "
            + f"/Contents {content_id} 0 R >>".encode("ascii")
        )
    kids = " ".join(f"{pid} 0 R" for pid in page_ids)
    objects[2] = f"<< /Type /Pages /Count {len(page_ids)} /Kids [{kids}] >>".encode("ascii")

    max_id = max(objects.keys())
    parts = [b"%PDF-1.4\n"]
    offsets = [0] * (max_id + 1)
    current_offset = len(parts[0])
    for obj_id in range(1, max_id + 1):
        offsets[obj_id] = current_offset
        block = f"{obj_id} 0 obj\n".encode("ascii") + objects[obj_id] + b"\nendobj\n"
        parts.append(block)
        current_offset += len(block)
    xref_offset = current_offset
    parts.append(f"xref\n0 {max_id + 1}\n".encode("ascii"))
    parts.append(b"0000000000 65535 f \n")
    for obj_id in range(1, max_id + 1):
        parts.append(f"{offsets[obj_id]:010d} 00000 n \n".encode("ascii"))
    parts.append(f"trailer\n<< /Size {max_id + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii"))
    return b"".join(parts)
