from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

def add_page_number(paragraph):
    """Safe page number field for DOCX (no errors)."""
    run = paragraph.add_run()

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)


def generate_docx(results, output_path, tone="Neutral"):
    doc = Document()

    # ---------------------------
    # HEADER
    # ---------------------------
    header = doc.sections[0].header
    h = header.paragraphs[0]
    h.text = "ClauseAI – Contract Risk Analysis Report"
    h.runs[0].font.bold = True
    h.runs[0].font.size = Pt(12)

    # ---------------------------
    # FOOTER
    # ---------------------------
    footer = doc.sections[0].footer
    f = footer.paragraphs[0]

    file_name = results.get("Contract Name", "Contract")
    f.add_run(f"{file_name} | Page ")
    add_page_number(f)

    # ---------------------------
    # TITLE
    # ---------------------------
    title = doc.add_heading("Contract Analysis Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run("Generated On: ").bold = True
    meta.add_run(datetime.now().strftime("%d-%m-%Y %I:%M %p"))
    meta.add_run("\nTone Applied: ").bold = True
    meta.add_run(tone)
    meta.add_run("\n\n")

    # ---------------------------
    # MAIN CONTENT
    # ---------------------------
    for key, value in results.items():
        doc.add_heading(str(key), level=2)

        # Convert any type to string
        value_str = str(value)

        # Tone styling
        if tone.lower() == "friendly":
            value_str = "😊 " + value_str
        elif tone.lower() == "strict":
            value_str = "⚠️ " + value_str

        p = doc.add_paragraph(value_str)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(10)

    # ---------------------------
    # SAVE DOCX
    # ---------------------------
    doc.save(output_path)
